"""
DOCX to PDF Conversion API Endpoint
Converts DOCX files to PDF format using python-docx and reportlab
"""

from http.server import BaseHTTPRequestHandler
import json
import base64
import io
from docx import Document
from reportlab.lib.pagesizes import letter
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle
from reportlab.lib import colors
from _auth import authenticate_request, check_payload_size, send_unauthorized, send_payload_too_large, get_cors_origin, send_cors_headers


class handler(BaseHTTPRequestHandler):
    def do_POST(self):
        # Auth check
        authed, result = authenticate_request(self)
        if not authed:
            send_unauthorized(self, result)
            return

        # Payload size check
        size_ok, size_result = check_payload_size(self)
        if not size_ok:
            send_payload_too_large(self)
            return

        origin = get_cors_origin(self)

        try:
            # Read request body
            content_length = int(self.headers['Content-Length'])
            post_data = self.rfile.read(content_length)

            # Parse JSON request
            request_data = json.loads(post_data.decode('utf-8'))

            # Extract base64 DOCX data
            if 'docx_base64' not in request_data:
                self.send_error(400, 'Missing docx_base64 field')
                return

            docx_base64 = request_data['docx_base64']

            # Decode base64 to bytes
            docx_bytes = base64.b64decode(docx_base64)

            # Create in-memory file objects
            docx_stream = io.BytesIO(docx_bytes)
            pdf_stream = io.BytesIO()

            # Parse DOCX
            doc = Document(docx_stream)

            # Create PDF
            pdf_doc = SimpleDocTemplate(
                pdf_stream,
                pagesize=letter,
                rightMargin=0.75*inch,
                leftMargin=0.75*inch,
                topMargin=0.75*inch,
                bottomMargin=0.75*inch
            )

            # Build content
            styles = getSampleStyleSheet()
            story = []

            # Process paragraphs
            for para in doc.paragraphs:
                if para.text.strip():
                    # Determine style based on paragraph formatting
                    if para.style.name.startswith('Heading'):
                        style = styles['Heading1']
                    else:
                        style = styles['Normal']

                    # Create paragraph
                    p = Paragraph(para.text, style)
                    story.append(p)
                    story.append(Spacer(1, 0.1*inch))

            # Process tables
            for table in doc.tables:
                table_data = []
                for row in table.rows:
                    row_data = [cell.text for cell in row.cells]
                    table_data.append(row_data)

                if table_data:
                    t = Table(table_data)
                    t.setStyle(TableStyle([
                        ('GRID', (0, 0), (-1, -1), 0.5, colors.grey),
                        ('BACKGROUND', (0, 0), (-1, 0), colors.lightgrey),
                        ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
                        ('VALIGN', (0, 0), (-1, -1), 'TOP'),
                        ('FONTSIZE', (0, 0), (-1, -1), 10),
                        ('BOTTOMPADDING', (0, 0), (-1, -1), 6),
                        ('TOPPADDING', (0, 0), (-1, -1), 6),
                    ]))
                    story.append(t)
                    story.append(Spacer(1, 0.2*inch))

            # Build PDF
            pdf_doc.build(story)

            # Get PDF bytes and encode to base64
            pdf_stream.seek(0)
            pdf_bytes = pdf_stream.read()
            pdf_base64 = base64.b64encode(pdf_bytes).decode('utf-8')

            # Send response
            self.send_response(200)
            self.send_header('Content-Type', 'application/json')
            send_cors_headers(self, origin)
            self.end_headers()

            response = {
                'success': True,
                'pdf_base64': pdf_base64,
                'message': 'DOCX successfully converted to PDF'
            }

            self.wfile.write(json.dumps(response).encode('utf-8'))

        except Exception as e:
            self.send_response(500)
            self.send_header('Content-Type', 'application/json')
            send_cors_headers(self, origin)
            self.end_headers()

            error_response = {
                'success': False,
                'error': str(e),
                'message': 'Failed to convert DOCX to PDF'
            }

            self.wfile.write(json.dumps(error_response).encode('utf-8'))

    def do_OPTIONS(self):
        origin = get_cors_origin(self)
        self.send_response(200)
        send_cors_headers(self, origin)
        self.send_header('Access-Control-Allow-Methods', 'POST, OPTIONS')
        self.send_header('Access-Control-Allow-Headers', 'Content-Type, Authorization')
        self.end_headers()
