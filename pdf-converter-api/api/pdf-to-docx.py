"""
PDF to DOCX Conversion API Endpoint
Extracts text from PDF and creates DOCX (lightweight approach for Vercel)
"""

from http.server import BaseHTTPRequestHandler
import json
import base64
import io
from PyPDF2 import PdfReader
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH


class handler(BaseHTTPRequestHandler):
    def do_POST(self):
        try:
            # Read request body
            content_length = int(self.headers['Content-Length'])
            post_data = self.rfile.read(content_length)

            # Parse JSON request
            request_data = json.loads(post_data.decode('utf-8'))

            # Extract base64 PDF data
            if 'pdf_base64' not in request_data:
                self.send_error(400, 'Missing pdf_base64 field')
                return

            pdf_base64 = request_data['pdf_base64']

            # Decode base64 to bytes
            pdf_bytes = base64.b64decode(pdf_base64)

            # Create in-memory file objects
            pdf_stream = io.BytesIO(pdf_bytes)

            # Read PDF using PyPDF2
            pdf_reader = PdfReader(pdf_stream)

            # Create new DOCX document
            doc = Document()

            # Set default font
            style = doc.styles['Normal']
            font = style.font
            font.name = 'Calibri'
            font.size = Pt(11)

            # Extract text from each page
            for page_num, page in enumerate(pdf_reader.pages):
                # Extract text
                text = page.extract_text()

                if text.strip():
                    # Add page content
                    if page_num > 0:
                        # Add page break for subsequent pages
                        doc.add_page_break()

                    # Split text into paragraphs (by double newline or single newline)
                    paragraphs = text.split('\n')

                    for para_text in paragraphs:
                        para_text = para_text.strip()
                        if para_text:
                            # Detect if it might be a heading (short, possibly all caps or title case)
                            is_heading = (
                                len(para_text) < 100 and
                                (para_text.isupper() or para_text.istitle()) and
                                not para_text.endswith('.')
                            )

                            if is_heading:
                                # Add as heading
                                para = doc.add_heading(para_text, level=2)
                            else:
                                # Add as normal paragraph
                                para = doc.add_paragraph(para_text)

            # Save DOCX to bytes
            docx_stream = io.BytesIO()
            doc.save(docx_stream)

            # Get DOCX bytes and encode to base64
            docx_stream.seek(0)
            docx_bytes = docx_stream.read()
            docx_base64 = base64.b64encode(docx_bytes).decode('utf-8')

            # Send response
            self.send_response(200)
            self.send_header('Content-Type', 'application/json')
            self.send_header('Access-Control-Allow-Origin', '*')
            self.end_headers()

            response = {
                'success': True,
                'docx_base64': docx_base64,
                'message': f'PDF successfully converted to DOCX ({len(pdf_reader.pages)} pages)'
            }

            self.wfile.write(json.dumps(response).encode('utf-8'))

        except Exception as e:
            self.send_response(500)
            self.send_header('Content-Type', 'application/json')
            self.send_header('Access-Control-Allow-Origin', '*')
            self.end_headers()

            error_response = {
                'success': False,
                'error': str(e),
                'message': 'Failed to convert PDF to DOCX'
            }

            self.wfile.write(json.dumps(error_response).encode('utf-8'))

    def do_OPTIONS(self):
        # Handle CORS preflight
        self.send_response(200)
        self.send_header('Access-Control-Allow-Origin', '*')
        self.send_header('Access-Control-Allow-Methods', 'POST, OPTIONS')
        self.send_header('Access-Control-Allow-Headers', 'Content-Type, x-vercel-protection-bypass, x-vercel-set-bypass-cookie')
        self.end_headers()
